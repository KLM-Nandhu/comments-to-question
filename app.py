import streamlit as st
from googleapiclient.discovery import build
from datetime import datetime
import os
import openai
import json
from docx import Document
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

# Get API keys from environment variables
YOUTUBE_API_KEY = os.environ.get("YOUTUBE_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

if not YOUTUBE_API_KEY:
    st.error("YouTube API key not found. Please set the YOUTUBE_API_KEY environment variable.")
    st.stop()

if not OPENAI_API_KEY:
    st.error("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
    st.stop()

youtube = build('youtube', 'v3', developerKey=YOUTUBE_API_KEY)
openai.api_key = OPENAI_API_KEY

def get_all_comments(video_id: str):
    try:
        comments = []
        nextPageToken = None
        while True:
            response = youtube.commentThreads().list(
                part='snippet',
                videoId=video_id,
                maxResults=100,
                pageToken=nextPageToken,
                order='time'
            ).execute()

            for item in response['items']:
                comment = item['snippet']['topLevelComment']['snippet']
                comments.append({
                    'author': comment['authorDisplayName'],
                    'text': comment['textDisplay'],
                    'likes': comment['likeCount'],
                    'published_at': datetime.strptime(comment['publishedAt'], "%Y-%m-%dT%H:%M:%SZ")
                })

            nextPageToken = response.get('nextPageToken')
            if not nextPageToken:
                break

        return comments
    except Exception as e:
        return f"An error occurred while fetching comments: {str(e)}"

def extract_questions(comments):
    comments_with_authors = [f"{comment['author']}: {comment['text']}" for comment in comments]
    all_comments_text = "\n".join(comments_with_authors)
    
    prompt = f"""Analyze the following YouTube comments and extract ALL direct and indirect questions about the video. Categorize them as either 'Direct' or 'Indirect'. Include the author of each question in parentheses.

Comments:
{all_comments_text}

Format your response as follows:
Direct Questions:
1. [Question 1] (Author1)
2. [Question 2] (Author2)
...

Indirect Questions:
1. [Question 1] (Author1)
2. [Question 2] (Author2)
...

If there are no questions in a category, write 'None found.' under that category.
"""

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that analyzes YouTube comments to extract questions."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000  # Increased token limit to accommodate more questions
    )

    return response.choices[0].message.content

def suggest_follow_up_questions(video_info, comments, extracted_questions):
    prompt = f"""Given the following information about a YouTube video and its comments, suggest 5 intelligent follow-up questions that users might ask to continue the conversation:

Video Title: {video_info['title']}
Video Views: {video_info['views']}
Video Likes: {video_info['likes']}
Number of Comments: {video_info['comments']}

Extracted Questions:
{extracted_questions}

Sample Comments:
{json.dumps(comments[:5], default=str)}

Suggest 5 follow-up questions that would encourage further discussion about the video content, address common themes in the comments, or explore aspects of the video that might not have been covered in the existing questions.
"""

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that suggests engaging follow-up questions based on YouTube video information and comments."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500
    )

    return response.choices[0].message.content

def get_video_info(video_id):
    try:
        response = youtube.videos().list(
            part='snippet,statistics',
            id=video_id
        ).execute()

        if 'items' in response:
            video = response['items'][0]
            return {
                'title': video['snippet']['title'],
                'views': video['statistics']['viewCount'],
                'likes': video['statistics']['likeCount'],
                'comments': video['statistics']['commentCount'],
                'published_at': video['snippet']['publishedAt']
            }
        else:
            return None
    except Exception as e:
        st.error(f"An error occurred while fetching video info: {str(e)}")
        return None

def analyze_sentiment(comments):
    positive_words = set(['good', 'great', 'awesome', 'excellent', 'amazing', 'love', 'like', 'best'])
    negative_words = set(['bad', 'terrible', 'awful', 'worst', 'hate', 'dislike', 'poor'])
    
    positive_count = 0
    negative_count = 0
    neutral_count = 0
    
    for comment in comments:
        text = comment['text'].lower()
        if any(word in text for word in positive_words):
            positive_count += 1
        elif any(word in text for word in negative_words):
            negative_count += 1
        else:
            neutral_count += 1
    
    total = positive_count + negative_count + neutral_count
    return {
        'positive': positive_count / total if total > 0 else 0,
        'negative': negative_count / total if total > 0 else 0,
        'neutral': neutral_count / total if total > 0 else 0
    }

def create_docx_report(video_info, comments, questions, sentiment, follow_up_questions):
    doc = Document()
    doc.add_heading('YouTube Video Analysis Report', 0)

    # Video Information
    doc.add_heading('Video Information', level=1)
    doc.add_paragraph(f"Title: {video_info['title']}")
    doc.add_paragraph(f"Views: {video_info['views']}")
    doc.add_paragraph(f"Likes: {video_info['likes']}")
    doc.add_paragraph(f"Comments: {video_info['comments']}")
    doc.add_paragraph(f"Published: {video_info['published_at']}")

    # Sentiment Analysis
    doc.add_heading('Sentiment Analysis', level=1)
    doc.add_paragraph(f"Positive: {sentiment['positive']:.2%}")
    doc.add_paragraph(f"Neutral: {sentiment['neutral']:.2%}")
    doc.add_paragraph(f"Negative: {sentiment['negative']:.2%}")

    # Extracted Questions
    doc.add_heading('Extracted Questions', level=1)
    doc.add_paragraph(questions)

    # Suggested Follow-up Questions
    doc.add_heading('Suggested Follow-up Questions', level=1)
    doc.add_paragraph(follow_up_questions)

    # Comments
    doc.add_heading('Comments', level=1)
    for comment in comments:
        doc.add_paragraph(f"Author: {comment['author']}")
        doc.add_paragraph(f"Text: {comment['text']}")
        doc.add_paragraph(f"Likes: {comment['likes']}")
        doc.add_paragraph(f"Published at: {comment['published_at']}")
        doc.add_paragraph("---")

    return doc

def analyze_comments(video_id):
    if video_id:
        with st.spinner("📊 Fetching and analyzing comments..."):
            comments = get_all_comments(video_id)
            if isinstance(comments, list):
                st.session_state.comments = comments
                st.session_state.comments.sort(key=lambda x: x['published_at'], reverse=True)
                st.session_state.questions = extract_questions(comments)
                st.session_state.video_info = get_video_info(video_id)
                st.session_state.sentiment = analyze_sentiment(comments)
                st.session_state.follow_up_questions = suggest_follow_up_questions(
                    st.session_state.video_info,
                    comments,
                    st.session_state.questions
                )
            else:
                st.error(comments)
    else:
        st.error("⚠️ Please enter a YouTube Video ID.")

st.set_page_config(layout="wide", page_title="Bent's Comment Analyzer 🎥💬")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600&display=swap');
    
    body {
        font-family: 'Poppins', sans-serif;
        background-color: #f0f2f5;
        color: #1a1a1a;
    }
    .main {
        max-width: 1200px;
        margin: 0 auto;
        background-color: #ffffff;
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    h1 {
        color: #3498db;
        text-align: center;
        font-size: 2.5em;
        margin-bottom: 1em;
    }
    h2 {
        color: #2c3e50;
        font-size: 1.8em;
        margin-top: 1em;
        margin-bottom: 0.5em;
    }
    .stTextInput>div>div>input {
        background-color: #f1f3f5;
        color: #333;
        border: 2px solid #3498db;
        border-radius: 5px;
        padding: 10px 15px;
    }
    .stButton>button {
        background-color: #3498db;
        color: white;
        border: none;
        padding: 10px 20px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        margin: 4px 2px;
        cursor: pointer;
        border-radius: 5px;
        transition: background-color 0.3s, color 0.3s;
    }
    .stButton>button:hover {
        background-color: #2980b9;
        color: #ffffff;
    }
    .small-button {
        padding: 5px 10px;
        font-size: 14px;
    }
    .comment {
        background-color: #f9f9f9;
        border-left: 4px solid #3498db;
        padding: 15px;
        margin-bottom: 15px;
        border-radius: 4px;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
    }
    .comment-author {
        font-weight: bold;
        color: #2c3e50;
    }
    .comment-date {
        font-size: 0.8em;
        color: #7f8c8d;
    }
    .comment-text {
        margin-top: 5px;
        color: #34495e;
    }
    .comment-likes {
        font-size: 0.9em;
        color: #e74c3c;
        margin-top: 5px;
    }
    .sort-button {
        background-color: #2ecc71;
        margin-bottom: 10px;
    }
    .sort-button:hover {
        background-color: #27ae60;
    }
</style>
""", unsafe_allow_html=True)

st.title("🎥 Bent's Comment Analyzer 💬")

video_id = st.text_input("🔍 Enter YouTube Video ID")

if 'comments' not in st.session_state:
    st.session_state.comments = []
if 'sort_order' not in st.session_state:
    st.session_state.sort_order = 'newest'
if 'show_comments' not in st.session_state:
    st.session_state.show_comments = 10
if 'questions' not in st.session_state:
    st.session_state.questions = None
if 'video_info' not in st.session_state:
    st.session_state.video_info = None
if 'sentiment' not in st.session_state:
    st.session_state.sentiment = None
if 'follow_up_questions' not in st.session_state:
    st.session_state.follow_up_questions = None

def toggle_sort_order():
    st.session_state.sort_order = 'oldest' if st.session_state.sort_order == 'newest' else 'newest'
    st.session_state.comments.sort(key=lambda x: x['published_at'], reverse=(st.session_state.sort_order == 'newest'))

def show_more_comments():
    st.session_state.show_comments = min(st.session_state.show_comments + 10, len(st.session_state.comments))

def show_less_comments():
    st.session_state.show_comments = max(st.session_state.show_comments - 10, 10)

if st.button("🚀 Analyze Comments", key="analyze_button"):
    analyze_comments(video_id)

# Display video information, sentiment analysis, and export data at the top
if st.session_state.video_info:
    st.markdown("## 📺 Video Information")
    st.markdown(f"**Title:** {st.session_state.video_info['title']}")
    st.markdown(f"**Views:** {st.session_state.video_info['views']}")
    st.markdown(f"**Likes:** {st.session_state.video_info['likes']}")
    st.markdown(f"**Comments:** {st.session_state.video_info['comments']}")
    st.markdown(f"**Published:** {st.session_state.video_info['published_at']}")

if st.session_state.sentiment:
    st.markdown("## 💭 Sentiment Analysis")
    col1, col2, col3 = st.columns(3)
    col1.metric("Positive", f"{st.session_state.sentiment['positive']:.2%}")
    col2.metric("Neutral", f"{st.session_state.sentiment['neutral']:.2%}")
    col3.metric("Negative", f"{st.session_state.sentiment['negative']:.2%}")

if st.session_state.comments:
    st.markdown("## 📤 Export Data")
    export_format = st.selectbox("Choose export format:", ["CSV", "JSON", "DOCX"])
    
    if st.button("Export Data"):
        if export_format == "CSV":
            csv = "Author,Text,Likes,Published At\n"
            for comment in st.session_state.comments:
                csv += f"{comment['author']},{comment['text']},{comment['likes']},{comment['published_at']}\n"
            st.download_button(
                label="Download CSV",
                data=csv,
                file_name="youtube_analysis.csv",
                mime="text/csv"
            )
        elif export_format == "JSON":
            data = {
                "video_info": st.session_state.video_info,
                "sentiment": st.session_state.sentiment,
                "questions": st.session_state.questions,
                "follow_up_questions": st.session_state.follow_up_questions,
                "comments": st.session_state.comments
            }
            json_str = json.dumps(data, default=str)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name="youtube_analysis.json",
                mime="application/json"
            )
        else:  # DOCX
            doc = create_docx_report(
                st.session_state.video_info,
                st.session_state.comments,
                st.session_state.questions,
                st.session_state.sentiment,
                st.session_state.follow_up_questions
            )
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                label="Download DOCX",
                data=bio.getvalue(),
                file_name="youtube_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Display comments and extracted questions
if st.session_state.comments:
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("<h2>📝 Comments</h2>", unsafe_allow_html=True)
        if st.button(f"{'🔽' if st.session_state.sort_order == 'newest' else '🔼'} Sort: {st.session_state.sort_order.capitalize()}", 
                    key="sort_button", help="Toggle between newest and oldest comments"):
            toggle_sort_order()
        
        for i, comment in enumerate(st.session_state.comments[:st.session_state.show_comments]):
            st.markdown(f"""
            <div class="comment">
                <div class="comment-author">👤 {comment['author']}</div>
                <div class="comment-date">🕒 {comment['published_at'].strftime('%Y-%m-%d %H:%M:%S')}</div>
                <div class="comment-text">{comment['text']}</div>
                <div class="comment-likes">❤️ {comment['likes']}</div>
            </div>
            """, unsafe_allow_html=True)
        
        col1_1, col1_2, col1_3 = st.columns([1,1,2])
        with col1_1:
            if st.session_state.show_comments < len(st.session_state.comments):
                if st.button("📥 Show More", key="show_more"):
                    show_more_comments()
        with col1_2:
            if st.session_state.show_comments > 10:
                if st.button("📤 Show Less", key="show_less"):
                    show_less_comments()
        with col1_3:
            st.write(f"Showing {st.session_state.show_comments} of {len(st.session_state.comments)} comments")
    
    with col2:
        st.markdown("<h2>❓ Extracted Questions</h2>", unsafe_allow_html=True)
        if st.session_state.questions:
            st.markdown(st.session_state.questions, unsafe_allow_html=True)
        else:
            st.info("🤔 No questions extracted yet. Try analyzing a video with more comments or discussions.")
        
        st.markdown("<h2>🔎 Suggested Follow-up Questions</h2>", unsafe_allow_html=True)
        if st.session_state.follow_up_questions:
            st.markdown(st.session_state.follow_up_questions, unsafe_allow_html=True)
        else:
            st.info("💡 Analyze a video to get suggested follow-up questions.")

st.markdown("---")
st.markdown("Developed with ❤️ using Streamlit")
